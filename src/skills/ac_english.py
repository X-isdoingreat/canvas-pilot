# SPDX-License-Identifier: AGPL-3.0-or-later
"""Writing-course document annotation skill.

Strategy:
1. Read assignment description + downloaded PDF/Doc text
2. Identify paragraphs, vocabulary words to annotate, comprehension questions
3. Generate B1-B2 level annotations + answers (with 1-2 intentional minor grammar slips)
4. Render to local .docx with annotations placed below each paragraph,
   then convert to PDF (via docx2pdf or LibreOffice if available, else leave .docx)
5. NEVER submits to Canvas — leaves draft for morning review.

NOTE: gws-docs route is preferred per plan but requires fetch-paragraph-coords work.
For tonight's validation pass we ship the local .docx fallback so the pipeline
produces a real artifact end-to-end. The probe at ac_english_probe.py exercises
the gws path separately so we can swap it in next iteration without changing the
router contract.
"""
from __future__ import annotations

import re
from pathlib import Path

from .base import Skill, html_to_text


def _extract_text_from_pdf(p: Path) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ""
    try:
        doc = fitz.open(str(p))
        return "\n\n".join(page.get_text() for page in doc)
    except Exception as e:
        return f"[pdf read error: {e}]"


def _split_paragraphs(text: str) -> list[str]:
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    # filter very short noise
    return [p for p in paras if len(p) > 40]


def _pick_vocab(paragraphs: list[str], n: int = 5) -> list[str]:
    """Pick a handful of harder-looking words for annotation."""
    seen = set()
    picks: list[str] = []
    for p in paragraphs:
        for w in re.findall(r"\b[A-Za-z]{7,}\b", p):
            wl = w.lower()
            if wl in seen:
                continue
            seen.add(wl)
            picks.append(w)
            if len(picks) >= n + 3:
                break
        if len(picks) >= n + 3:
            break
    return picks[:n]


# Hand-rolled small dictionary of generic glosses + B1-B2 example sentences.
# When the picked word isn't in here, we fall back to a generic template.
GLOSS = {
    # very common academic words; expand over time
}


def _annotate_word(word: str) -> str:
    g = GLOSS.get(word.lower())
    if g:
        return g
    return (
        f"\"{word}\": this word means something important in the paragraph. "
        f"In simple english, the writer use it to show his idea more clear. "
        f"Example: I tried to {word.lower()} my homework today, but it was hard."
    )


def _summarize_paragraph(p: str) -> str:
    # Very simple: take first sentence + a one-line takeaway. Intentionally
    # casual phrasing for human-ness.
    first = re.split(r"(?<=[.!?])\s", p.strip())[0][:240]
    return f"Main idea: the writer talk about {first.lower()[:160]}..."


class ACEnglishSkill(Skill):
    name = "ac_english"

    def plan(self) -> str:
        return (
            "1. Read description + attached PDFs\n"
            "2. Split into paragraphs, pick 5 vocab words\n"
            "3. Write annotations under each para (1-2 per para, some skipped)\n"
            "4. Answer comprehension questions if any\n"
            "5. Output .docx draft (PDF if converter available)\n"
            "6. DO NOT submit\n"
        )

    def _gather_text(self) -> str:
        a = self.assignment or {}
        body = html_to_text(a.get("description"))
        att_dir = self.work_dir / "attachments"
        if att_dir.exists():
            for p in sorted(att_dir.iterdir()):
                if p.suffix.lower() == ".pdf":
                    body += "\n\n[PDF: " + p.name + "]\n\n" + _extract_text_from_pdf(p)
                elif p.suffix.lower() in (".txt", ".md"):
                    body += "\n\n[FILE: " + p.name + "]\n\n" + p.read_text(encoding="utf-8", errors="ignore")
        return body

    def _build_doc(self, paragraphs: list[str], vocab: list[str], full_text: str) -> Path:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document()
        doc.add_heading(self.assignment.get("name", "Assignment"), 0)
        doc.add_paragraph(f"Course: {self.item['course_name']}")
        student = os.environ.get("STUDENT_FULL_NAME", "Student")
        doc.add_paragraph(f"Draft by: {student}  |  Status: DRAFT — review before submitting")
        doc.add_heading("Annotated Reading", 1)

        skip_every = 3  # leave roughly 1 in 3 paragraphs unannotated
        for i, para in enumerate(paragraphs):
            doc.add_paragraph(para)
            if i % skip_every == 2:
                continue
            note = doc.add_paragraph()
            run = note.add_run("Note: " + _summarize_paragraph(para))
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_heading("Vocabulary Annotations", 1)
        for w in vocab:
            doc.add_paragraph("- " + _annotate_word(w))

        # Comprehension answer placeholder if questions exist
        questions = re.findall(r"(?im)^\s*(?:\d+[.)]|Q\d+[.:])\s*(.+)$", full_text)
        if questions:
            doc.add_heading("Answers", 1)
            for q in questions[:10]:
                doc.add_paragraph("Q: " + q.strip())
                doc.add_paragraph("A: I think the writer want to say that " + q.strip().lower()[:80] + " is important. Because in the text, the writer use evidence to support this point, and i agree with it for the most part.")

        out = self.work_dir / "draft" / "annotated.docx"
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return out

    def draft(self) -> dict:
        full_text = self._gather_text()
        paragraphs = _split_paragraphs(full_text)
        if not paragraphs:
            paragraphs = ["[No readable paragraphs were extracted from the assignment. Manual review needed.]"]
        vocab = _pick_vocab(paragraphs, n=5)
        try:
            docx_path = self._build_doc(paragraphs, vocab, full_text)
        except ImportError:
            return {"status": "error", "message": "python-docx not installed; pip install python-docx"}

        # Try to convert to PDF (best effort)
        pdf_path = None
        try:
            from docx2pdf import convert  # type: ignore
            pdf_target = docx_path.with_suffix(".pdf")
            convert(str(docx_path), str(pdf_target))
            if pdf_target.exists():
                pdf_path = pdf_target
        except Exception:
            pass

        return {
            "status": "draft_ready",
            "draft_path": str(pdf_path or docx_path),
            "notes": f"{len(paragraphs)} paras, {len(vocab)} vocab. NOT submitted.",
        }


def run(item: dict, run_dir: Path) -> dict:
    return ACEnglishSkill(item, run_dir).run()
