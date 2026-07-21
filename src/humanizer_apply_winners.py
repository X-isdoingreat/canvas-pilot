# SPDX-License-Identifier: AGPL-3.0-or-later
"""Apply per-segment winners (from a winners JSON) to the input docx, write output
docx + humanizer_log.json per canvas-humanizer §8 + §10.

Input: winners.json with shape:
  {
    "segments": [
      {"seg_id": "S0", "doc_paragraph_index": 5, "intra_para_index": 0,
       "original": "...", "final_text": "...", "strategy": "...", "method": "...",
       "divergence": 0.46, "candidates_considered": 6, ...}
    ]
  }

Reads _state.json from the same dir for paragraph passthroughs + lock substitutions.
Writes humanized.docx and humanizer_log.json.
"""
import argparse
import json
import re
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document


def banned_words_check(text: str, banned: list) -> list:
    """Return banned words found in text (case-insensitive, whole-word)."""
    found = []
    for w in banned:
        if re.search(r'\b' + re.escape(w) + r'\b', text, re.I):
            found.append(w)
    return found


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--state", required=True)
    ap.add_argument("--winners", required=True)
    ap.add_argument("--scored", required=True)
    ap.add_argument("--out-docx", required=True)
    ap.add_argument("--out-log", required=True)
    args = ap.parse_args()

    state = json.loads(Path(args.state).read_text(encoding="utf-8"))
    winners = json.loads(Path(args.winners).read_text(encoding="utf-8"))
    scored = json.loads(Path(args.scored).read_text(encoding="utf-8"))

    # Map seg_id → winner
    winner_by_seg = {w["seg_id"]: w for w in winners["segments"]}

    # Map seg_id → scored block (for candidate counts)
    scored_by_seg = {s["seg_id"]: s for s in scored["segments"]}

    BANNED = [
        "delve","leverage","tapestry","intricate","multifaceted","plethora",
        "paradigm","holistic","harness","nuance","interplay","landscape",
        "realm","journey","robust","seamless","dynamic","Moreover","Furthermore",
        "In conclusion","In summary","To summarize","It's worth noting",
        "It is important to note","In today's world","Navigate","Embark",
        "henceforth","notwithstanding","vis-à-vis",
    ]

    # Reassemble paragraph by paragraph (§8)
    final_paragraphs = {}
    para_drifts = []
    seg_log = []

    for p in state["humanizable_paragraphs"]:
        pi = p["doc_paragraph_index"]
        # All sentence final texts in intra_para_index order
        sent_finals = []
        for s in p["sentences"]:
            if s["humanizable"] and s["seg_id"] in winner_by_seg:
                final_text = winner_by_seg[s["seg_id"]]["final_text"]
            else:
                final_text = s["original"]
            sent_finals.append((s["intra_para_index"], final_text))
        sent_finals.sort(key=lambda t: t[0])
        para_text = " ".join(t[1] for t in sent_finals)
        final_paragraphs[pi] = para_text

        # Drift check
        orig_wc = p["paragraph_word_count"]
        final_wc = len(para_text.split())
        ratio = final_wc / orig_wc if orig_wc else 0.0
        drift = {"pi": pi, "orig_wc": orig_wc, "final_wc": final_wc, "ratio": round(ratio, 4)}
        if not (0.85 <= ratio <= 1.15):
            drift["warning"] = f"outside ±15% band"
        para_drifts.append(drift)

    # Compose seg_log entries
    total_segments = 0
    humanizable_segments = 0
    r_wins = 0
    p_wins = 0
    div_values = []
    fallback_count = 0

    for p in state["humanizable_paragraphs"]:
        for s in p["sentences"]:
            total_segments += 1
            if s["humanizable"]:
                humanizable_segments += 1
                if s["seg_id"] in winner_by_seg:
                    w = winner_by_seg[s["seg_id"]]
                    sc = scored_by_seg.get(s["seg_id"], {})
                    candidates_considered = len(sc.get("candidates", []))
                    candidates_passed_meaning = sum(1 for c in sc.get("candidates", [])
                                                    if c.get("lock_gate_pass") and c.get("wc_gate_pass"))
                    candidates_lost_lock = sum(1 for c in sc.get("candidates", [])
                                               if not c.get("lock_gate_pass", True))
                    candidates_wc_failed = sum(1 for c in sc.get("candidates", [])
                                               if not c.get("wc_gate_pass", True))
                    banned_found = banned_words_check(w["final_text"], BANNED)
                    final_wc = len(w["final_text"].split())
                    if w.get("strategy") == "roundtrip":
                        r_wins += 1
                    elif w.get("strategy") == "paraphrase":
                        p_wins += 1
                    div_values.append(w["divergence"])
                    seg_log.append({
                        "seg_id": s["seg_id"],
                        "doc_paragraph_index": s["doc_paragraph_index"],
                        "intra_para_index": s["intra_para_index"],
                        "original_word_count": s["word_count"],
                        "final_word_count": final_wc,
                        "winning_strategy": w.get("strategy"),
                        "winning_method": w.get("method"),
                        "candidates_considered": candidates_considered,
                        "candidates_passed_meaning_gate": candidates_passed_meaning,
                        "candidates_lost_lock": candidates_lost_lock,
                        "candidates_wc_gate_failed": candidates_wc_failed,
                        "final_divergence": w["divergence"],
                        "banned_words_found": banned_found,
                        "status": "ok" if not banned_found else "ok_with_banned_words",
                        "fallback_to_original": False,
                    })
                else:
                    fallback_count += 1
                    seg_log.append({
                        "seg_id": s["seg_id"],
                        "doc_paragraph_index": s["doc_paragraph_index"],
                        "intra_para_index": s["intra_para_index"],
                        "original_word_count": s["word_count"],
                        "final_word_count": s["word_count"],
                        "status": "fallback_to_original",
                        "fallback_to_original": True,
                    })

    avg_div = round(sum(div_values) / len(div_values), 4) if div_values else 0.0

    # Write docx (§8a)
    doc = Document(state["draft_path"])
    for i, para in enumerate(doc.paragraphs):
        if i in final_paragraphs:
            new_text = final_paragraphs[i]
            if new_text != para.text:
                if para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run(new_text)
    doc.save(args.out_docx)

    # Write log (§10)
    log = {
        "version": 2,
        "draft_path": state["draft_path"],
        "output_path": args.out_docx,
        "voice_register": state["voice_register"],
        "student_identity": state["student_identity"],
        "config": {
            "K_total": 6,
            "K_roundtrip": 3,
            "K_paraphrase": 3,
            "languages_used": ["zh", "ja", "de"],
            "paraphrase_strategies_used": ["clause_reorder", "voice_flip", "nominalize_flip"],
            "segment_level": "sentence",
            "word_count_tolerance_per_segment": "sliding (±20-40%)",
            "word_count_tolerance_per_paragraph": 0.15,
            "detector_api": "manual",
            "detector_target": 30,
            "parallel_candidate_dispatch": False,
            "dispatch_mode": "inline_subagent",
            "run_label": "variance_test_run_3",
        },
        "total_paragraphs": 9,
        "total_segments": total_segments,
        "humanizable_segments": humanizable_segments,
        "locked_segments": 5,  # name-block + title
        "avg_divergence": avg_div,
        "fallback_count": fallback_count,
        "r_wins": r_wins,
        "p_wins": p_wins,
        "status": "ok" if fallback_count == 0 else "partial",
        "paragraph_word_count_drifts": para_drifts,
        "segments": seg_log,
    }
    Path(args.out_log).write_text(json.dumps(log, indent=2, ensure_ascii=False), encoding="utf-8")

    print(f"status: {log['status']}")
    print(f"total_segments: {total_segments}")
    print(f"humanizable_segments: {humanizable_segments}")
    print(f"r_wins: {r_wins}")
    print(f"p_wins: {p_wins}")
    print(f"avg_divergence: {avg_div}")
    print(f"fallback_count: {fallback_count}")
    # winning_method distribution
    from collections import Counter
    methods = Counter(w.get("winning_method") for w in seg_log if w.get("winning_method"))
    print(f"winning_method distribution: {dict(methods)}")
    print(f"docx: {args.out_docx}")
    print(f"log: {args.out_log}")


if __name__ == "__main__":
    main()
