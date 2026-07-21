# SPDX-License-Identifier: AGPL-3.0-or-later
from __future__ import annotations

import importlib
import json
import re
from pathlib import Path

import fitz
import pytest
from docx import Document

from src.ac_eng_verify import (
    CONTENT_HIGHLIGHT,
    CONTENT_NOTE_TEXT,
    VOCAB_DEF_TEXT,
    VOCAB_HIGHLIGHT,
    verify_ac_eng_draft,
)
from src.awkward_transformation_assigner import (
    assign_transformations,
    contains_banned,
    validate_transformation,
)
from src.course_artifacts import (
    ensure_stable_work_dir,
    html_to_text,
    redact_behavioral_rules,
    stable_work_dir,
    unresolved_placeholders,
    write_course_result,
)
from src.humanizer_score import score_candidate
from src.humanizer_segment_extract import extract_locks
from src.humanizer_segmentation import split_sentences
from src.run_state import validate_result
from src.zybooks_spec_parser import parse_homework_spec


ROOT = Path(__file__).resolve().parents[1]
SKILL_NAMES = (
    "canvas-generic",
    "canvas-essay",
    "canvas-reading-annotation",
    "canvas-zybooks",
    "canvas-humanizer",
    "canvas-humanizer-loop",
    "canvas-humanizer-surgical",
    "canvas-awkward-syntax",
)
COURSE_SKILLS = SKILL_NAMES[:4]
CANONICAL_RESULT_STATUSES = {"draft_ready", "submitted", "skipped", "error"}


def skill_text(name: str) -> str:
    path = ROOT / ".agents" / "skills" / name / "SKILL.md"
    assert path.is_file()
    return path.read_text(encoding="utf-8")


def test_skill_family_is_codex_native_and_compact() -> None:
    stale = re.compile(
        r"(?:\.claude[/\\]|\bclaude\b|agent tool|skill tool|subagent_type|"
        r"auto[- ]ported|cc context)",
        re.IGNORECASE,
    )
    for name in SKILL_NAMES:
        text = skill_text(name)
        assert not stale.search(text), name
        assert "native Codex" in text, name
        assert len(text.splitlines()) < 500, name


def test_stable_id_work_directory_is_explicit_everywhere() -> None:
    expected = "course-<course_id>__assignment-<assignment_id>"
    for name in SKILL_NAMES:
        text = skill_text(name)
        assert expected in text, name
        assert "src.course_artifacts" in text, name


def test_course_skills_use_only_canonical_result_statuses() -> None:
    for name in COURSE_SKILLS:
        text = skill_text(name)
        assert "draft_ready" in text
        assert "skipped" in text
        assert "error" in text
        found = set(
            re.findall(
                r'status\s*=\s*"([A-Za-z_]+)',
                text,
            )
        )
        assert found <= CANONICAL_RESULT_STATUSES, (name, found)


def test_course_skills_cannot_directly_mutate_canvas_or_zybooks() -> None:
    forbidden_calls = re.compile(
        r"(?:\bcv\.(?:post|put)\s*\(|\bsubmit_(?:files|text|url)\s*\(|"
        r"\bzybooks_client\.post\s*\()"
    )
    for name in COURSE_SKILLS:
        text = skill_text(name)
        assert not forbidden_calls.search(text), name
        assert "signed" in text.lower() and "receipt" in text.lower(), name


@pytest.mark.parametrize(
    ("module_name", "symbols"),
    (
        (
            "src.course_artifacts",
            (
                "stable_work_dir",
                "ensure_stable_work_dir",
                "atomic_write_text",
                "redact_behavioral_rules",
                "unresolved_placeholders",
                "write_course_result",
            ),
        ),
        ("src.ac_eng_router", ("route_ac_eng_assignment",)),
        ("src.ac_eng_verify", ("verify_ac_eng_draft",)),
        ("src.overlay_utils", ("canvas_generic_overlay_path",)),
        ("src.humanizer_segmentation", ("split_sentences", "paragraph_segment_counts")),
        ("src.humanizer_segment_extract", ("extract_locks",)),
        ("src.humanizer_score", ("score_candidate", "word_count_tolerance")),
        (
            "src.awkward_transformation_assigner",
            (
                "assign_transformations",
                "contains_banned",
                "validate_transformation",
                "write_assignment_report",
            ),
        ),
        ("src.zybooks_spec_parser", ("parse_homework_spec",)),
        (
            "src.zybooks_client",
            ("whoami", "exercises_for_section", "exercise_to_dict"),
        ),
    ),
)
def test_referenced_helper_symbols_exist(module_name: str, symbols: tuple[str, ...]) -> None:
    module = importlib.import_module(module_name)
    missing = [symbol for symbol in symbols if not hasattr(module, symbol)]
    assert not missing, (module_name, missing)


def test_stable_work_dir_and_result_contract_fixture(tmp_path: Path) -> None:
    run_dir = tmp_path / "runs" / "2026-07-18"
    expected = run_dir / "course-course_7__assignment-assignment_19"
    assert stable_work_dir(run_dir, "course_7", "assignment_19") == expected
    work = ensure_stable_work_dir(run_dir, "course_7", "assignment_19")
    draft = work / "draft" / "answer.md"
    draft.parent.mkdir(parents=True)
    draft.write_text("A source-grounded local draft.\n", encoding="utf-8")
    (work / "verification.log").write_text(
        "PASS | source grounding | measured: complete\n", encoding="utf-8"
    )

    result_path = write_course_result(
        work,
        status="draft_ready",
        draft_path=draft,
        notes="local fixture",
        metadata={"skill": "canvas-generic"},
    )
    result = json.loads(result_path.read_text(encoding="utf-8"))
    assert validate_result(result, work_dir=work)["status"] == "draft_ready"
    assert Path(result["draft_path"]) == draft


def test_generic_and_essay_local_artifact_fixtures(tmp_path: Path) -> None:
    assert html_to_text("<p>Write <b>600 words</b>.</p>") == "Write 600 words."
    redacted = redact_behavioral_rules(
        "Write 600 words. Do not use AI. Cite two sources."
    )
    assert "Write 600 words" in redacted
    assert "Cite two sources" in redacted
    assert "use AI" not in redacted
    assert unresolved_placeholders("Answer [answer needed] and [TODO].") == [
        "[TODO]",
        "[answer needed]",
    ]

    essay = tmp_path / "essay.docx"
    document = Document()
    document.add_heading("Synthetic Essay", level=1)
    document.add_paragraph("A concrete claim grounded in a synthetic source.")
    document.save(essay)
    reopened = Document(essay)
    assert "synthetic source" in "\n".join(p.text for p in reopened.paragraphs)


def _build_annotation_fixture(original_path: Path, draft_path: Path) -> None:
    original = fitz.open()
    page = original.new_page(width=612, height=792)
    page.insert_text((90, 80), "1")
    page.insert_text((90, 100), "Alpha concept explains a synthetic paragraph.")
    page.insert_text((90, 160), "________________________________________")
    original.save(original_path)
    original.close()

    draft = fitz.open(original_path)
    page = draft[0]
    alpha_rect = page.search_for("Alpha")[0]
    vocab = page.add_highlight_annot(alpha_rect)
    vocab.set_colors(stroke=VOCAB_HIGHLIGHT)
    vocab.update()
    content_rect = page.search_for("paragraph")[0]
    content = page.add_highlight_annot(content_rect)
    content.set_colors(stroke=CONTENT_HIGHLIGHT)
    content.update()

    for index in range(5):
        page.insert_text(
            (20, 220 + index * 12),
            f"term{index}",
            fontsize=6.5,
            color=VOCAB_DEF_TEXT,
        )
    page.insert_text(
        (90, 130),
        "one concrete content note",
        fontsize=7.5,
        color=CONTENT_NOTE_TEXT,
    )
    underscore_rects = page.search_for("_")
    line_width = max(rect.x1 for rect in underscore_rects) - min(
        rect.x0 for rect in underscore_rects
    )
    answer = "This synthetic answer gives concrete grounded detail"
    while fitz.get_text_length(answer + " detail", fontname="helv", fontsize=10) <= 0.97 * line_width:
        answer += " detail"
    assert 0.92 * line_width <= fitz.get_text_length(answer, fontname="helv", fontsize=10) <= line_width
    page.insert_text((90, 160), answer, fontsize=10, color=(0.0, 0.15, 0.80))
    draft.save(draft_path)
    draft.close()


def test_annotation_pipeline_fixture_passes_real_verifier(tmp_path: Path) -> None:
    original = tmp_path / "reading.pdf"
    draft = tmp_path / "reading-annotated.pdf"
    _build_annotation_fixture(original, draft)
    report = verify_ac_eng_draft(draft, original)
    assert report["all_passed"], report["log_text"]


def test_zybooks_parser_fixture_extracts_only_graded_column() -> None:
    html = """
    <table>
      <tr><td>Suggested Practice</td><td>Graded for Honest Effort</td></tr>
      <tr><td>1) 1.6.1 a, c</td><td>1) 1.7.3</td></tr>
      <tr><td>2) 1.6.2 a</td><td>2) 1.7.7 b, c, i</td></tr>
    </table>
    """
    assert parse_homework_spec(html) == [
        (1, 7, 3, None),
        (1, 7, 7, ["b", "c", "i"]),
    ]


def test_humanizer_and_awkward_pipeline_fixtures() -> None:
    sentences = split_sentences(
        'Dr. Example wrote "Synthetic Claim" in 2024. The result was clear.'
    )
    assert len(sentences) == 2
    locks = extract_locks(sentences[0], ["Dr. Example"])
    assert {lock["text"] for lock in locks} >= {"Dr. Example", '"Synthetic Claim"', "2024"}

    scored = score_candidate(
        original="The result was clear for the entire group.",
        masked_original="The result was clear for the entire group.",
        original_wc=8,
        candidate="For the entire group, the result remained clear.",
        lock_substitutions=[],
    )
    assert scored["lock_gate_pass"]
    assert scored["wc_gate_pass"]
    assert scored["divergence"] > 0

    records = [
        {
            "seg_id": "S0",
            "doc_paragraph_index": 0,
            "intra_para_index": 0,
            "role": "intro_opener",
            "v0_text": "A synthetic opening sentence contains enough words for transformation.",
        },
        {
            "seg_id": "S1",
            "doc_paragraph_index": 1,
            "intra_para_index": 0,
            "role": "body_elaboration",
            "v0_text": "This synthetic body sentence explains the evidence with concrete detail.",
        },
    ]
    assigned = assign_transformations(records)
    assert all(record.get("transformation") for record in assigned)
    assert validate_transformation(
        "The data show a synthetic result.",
        "The fact that the data show a synthetic result is clear.",
        "fact_noun_complement",
    )["ok"]
    assert not validate_transformation(
        "The data show a synthetic result.",
        "The data show a synthetic result: clearly.",
        "minimal_lexical",
    )["ok"]
    assert not contains_banned("This is a concrete sentence.")
    assert contains_banned("Moreover, this is generic.") == "Moreover"
